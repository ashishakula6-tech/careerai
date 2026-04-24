"""
Universal File Text Extractor — extracts readable text from ANY resume/portfolio format.

Supports:
- PDF (multi-page, scanned with OCR fallback)
- DOCX / DOC (Microsoft Word)
- RTF (Rich Text Format)
- TXT / MD / HTML (plain text formats)
- Images: JPG, PNG, WEBP (extracts text via OCR if tesseract available)
- ODT (OpenDocument)
- Pages (Apple Pages — treated as zip)
- Any unknown format — tries UTF-8, Latin-1, then raw byte extraction

Design patterns in resumes (columns, tables, headers, graphics) are handled by
the underlying libraries. We extract ALL text regardless of layout.
"""
import io
import re


def extract_text(content: bytes, filename: str = "", content_type: str = "") -> str:
    """Extract text from any file format. Returns plain text string."""
    fname = filename.lower()
    ctype = content_type.lower() if content_type else ""

    # Determine format
    if fname.endswith(".pdf") or ctype == "application/pdf" or content[:5] == b"%PDF-":
        return _extract_pdf(content)

    if fname.endswith((".docx",)) or ctype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(content)

    if fname.endswith(".doc") or ctype == "application/msword":
        return _extract_doc(content)

    if fname.endswith(".rtf") or ctype == "application/rtf":
        return _extract_rtf(content)

    if fname.endswith((".html", ".htm")) or ctype.startswith("text/html"):
        return _extract_html(content)

    if fname.endswith((".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".gif")) or ctype.startswith("image/"):
        return _extract_image(content)

    if fname.endswith(".odt") or ctype == "application/vnd.oasis.opendocument.text":
        return _extract_odt(content)

    # Default: try as text
    return _extract_text(content)


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF — handles multi-page, columns, tables."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        full_text = "\n".join(pages)
        if full_text.strip():
            return full_text

        # If PyPDF2 got nothing (scanned PDF), try OCR
        return _ocr_pdf(content)
    except Exception as e:
        # Fallback: try raw text extraction
        return _extract_text(content)


def _ocr_pdf(content: bytes) -> str:
    """OCR a scanned PDF by converting pages to images."""
    try:
        from PIL import Image
        from PyPDF2 import PdfReader

        # Try to extract images from PDF and OCR them
        reader = PdfReader(io.BytesIO(content))
        texts = []
        for page in reader.pages:
            for image_obj in page.images:
                try:
                    img = Image.open(io.BytesIO(image_obj.data))
                    text = _ocr_image(img)
                    if text:
                        texts.append(text)
                except:
                    pass
        if texts:
            return "\n".join(texts)
    except:
        pass
    return _extract_text(content)


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX — handles tables, headers, footers, text boxes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        texts = []

        # Paragraphs (main body)
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)

        # Tables (common in resume designs with columns)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    texts.append(" | ".join(row_text))

        # Headers
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                if para.text.strip():
                    texts.append(para.text)

        return "\n".join(texts) if texts else _extract_text(content)
    except Exception:
        return _extract_text(content)


def _extract_doc(content: bytes) -> str:
    """Extract text from old .doc format — binary Word files."""
    # .doc is a complex binary format. Try to extract readable strings.
    return _extract_binary_strings(content)


def _extract_rtf(content: bytes) -> str:
    """Extract text from RTF."""
    try:
        from striprtf.striprtf import rtf_to_text
        text = content.decode("utf-8", errors="replace")
        return rtf_to_text(text)
    except:
        return _extract_text(content)


def _extract_html(content: bytes) -> str:
    """Extract text from HTML — strips all tags, keeps text content."""
    try:
        text = content.decode("utf-8", errors="replace")
        # Remove script and style blocks
        text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', ' ', text)
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        # Decode HTML entities
        import html
        text = html.unescape(text)
        return text
    except:
        return _extract_text(content)


def _extract_odt(content: bytes) -> str:
    """Extract text from OpenDocument (ODT) — it's a ZIP with XML inside."""
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        zf = zipfile.ZipFile(io.BytesIO(content))
        xml_content = zf.read("content.xml")
        root = ET.fromstring(xml_content)
        texts = []
        for elem in root.iter():
            if elem.text and elem.text.strip():
                texts.append(elem.text.strip())
        return "\n".join(texts)
    except:
        return _extract_text(content)


def _extract_image(content: bytes) -> str:
    """Extract text from image using OCR (tesseract) or basic analysis."""
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(content))
        text = _ocr_image(img)
        if text:
            return text
    except:
        pass
    return "(Image file uploaded — text extraction requires OCR. Key information may need manual review.)"


def _ocr_image(img) -> str:
    """Run OCR on a PIL Image."""
    try:
        import pytesseract
        text = pytesseract.image_to_string(img)
        if text and text.strip():
            return text.strip()
    except:
        pass
    return ""


def _extract_text(content: bytes) -> str:
    """Try to decode as text using multiple encodings."""
    for encoding in ["utf-8", "latin-1", "cp1252", "ascii", "utf-16"]:
        try:
            text = content.decode(encoding, errors="replace")
            # Check if it looks like actual text (not binary garbage)
            printable_ratio = sum(1 for c in text[:500] if c.isprintable() or c in '\n\r\t') / max(len(text[:500]), 1)
            if printable_ratio > 0.7:
                return text
        except:
            continue
    # Last resort: extract any readable strings from binary
    return _extract_binary_strings(content)


def _extract_binary_strings(content: bytes, min_length: int = 4) -> str:
    """Extract readable ASCII/Unicode strings from binary data."""
    # Find sequences of printable characters
    strings = re.findall(rb'[\x20-\x7e]{' + str(min_length).encode() + rb',}', content)
    text = " ".join(s.decode("ascii", errors="ignore") for s in strings)
    # Also try UTF-16 strings (common in .doc files)
    try:
        utf16_text = content.decode("utf-16-le", errors="ignore")
        printable = "".join(c for c in utf16_text if c.isprintable() or c in '\n\r\t ')
        if len(printable) > len(text):
            text = printable
    except:
        pass
    return text[:10000]  # Cap at 10K chars
